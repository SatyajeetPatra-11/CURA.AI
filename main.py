import streamlit as st
from langchain_community.llms import Ollama
from langchain.prompts import ChatPromptTemplate
from langchain.schema import StrOutputParser
from langchain.schema.runnable import Runnable
import asyncio
import concurrent.futures
from streamlit_option_menu import option_menu
from pymed import PubMed
from typing import List
from haystack import component, Document, Pipeline
from haystack.components.generators import HuggingFaceTGIGenerator
from haystack.components.builders.prompt_builder import PromptBuilder
from dotenv import load_dotenv
import os
from crewai import Agent, Task, Crew
from crewai_tools import ScrapeWebsiteTool, SerperDevTool
from langchain_openai import ChatOpenAI
#from docx import Document

from io import BytesIO
import base64
from llama_parse import LlamaParse
import sqlite3
import requests
import firebase_admin
from firebase_admin import credentials
import json

# Load environment variables
load_dotenv()

# Set page configuration
st.set_page_config(page_title="CURA.AI", layout="wide")

# Custom CSS for styling
st.markdown("""
    <style>
    body {
        background-color: #f8f9fa;
    }
    .sidebar .sidebar-content {
        background-color: #343a40;
        color: white;
    }
    .sidebar .sidebar-content .option-menu .nav-link {
        color: white;
    }
    .sidebar .sidebar-content .option-menu .nav-link:hover {
        background-color: #495057;
        color: white;
    }
    .sidebar .sidebar-content .option-menu .nav-link-selected {
        background-color: #17a2b8;
        color: white;
    }
    .main .block-container {
        padding: 2rem;
        background-color: white;
        border-radius: 8px;
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
    }
    .stButton button {
        background-color: #02ab21;
        color: white;
        border-radius: 5px;
        padding: 0.5rem 1rem;
        border: none;
        cursor: pointer;
        transition: background-color 0.3s, color 0.3s; /* Smooth transition for hover effects */
    }
    .stButton button:hover {
        background-color: #018f10; /* New background color on hover */
        color: #e0e0e0; /* New text color on hover */
    }
    .stTextInput>div>input {
        border-radius: 5px;
        border: 1px solid #ced4da;
        padding: 0.5rem;
    }
    .stTextArea textarea {
        border-radius: 5px;
        border: 1px solid #ced4da;
        padding: 0.5rem;
    }
    </style>
""", unsafe_allow_html=True)

# Google Analytics setup (if needed)
st.markdown(
    f"""
    <!-- Global site tag (gtag.js) - Google Analytics -->
    <script async src="https://www.googletagmanager.com/gtag/js?id={os.getenv('analytics_tag')}"></script>
    <script>
        window.dataLayer = window.dataLayer || [];
        function gtag(){{dataLayer.push(arguments);}}
        gtag('js', new Date());
        gtag('config', '{os.getenv('analytics_tag')}');
    </script>
    """,
    unsafe_allow_html=True
)

# Initialize PubMed client
pubmed = PubMed(tool="Haystack2.0Prototype", email="dummyemail@gmail.com")

# Firebase API key and URLs
FIREBASE_API_KEY = "AIzaSyApr-etDzcGcsVcmaw7R7rPxx3A09as7uw"
SIGN_UP_URL = "https://identitytoolkit.googleapis.com/v1/accounts:signUp"
SIGN_IN_URL = "https://identitytoolkit.googleapis.com/v1/accounts:signInWithPassword"
RESET_PASSWORD_URL = "https://identitytoolkit.googleapis.com/v1/accounts:sendOobCode"

def sign_up(email, password, username=None):
    try:
        payload = {
            "email": email,
            "password": password,
            "returnSecureToken": True
        }
        if username:
            payload["displayName"] = username
        response = requests.post(SIGN_UP_URL, params={"key": FIREBASE_API_KEY}, json=payload)
        response_data = response.json()
        if 'email' in response_data:
            return response_data['email']
        else:
            st.warning(response_data.get('error', {}).get('message', 'Signup failed'))
    except Exception as e:
        st.warning(f'Signup failed: {e}')

def sign_in(email, password):
    try:
        payload = {
            "email": email,
            "password": password,
            "returnSecureToken": True
        }
        response = requests.post(SIGN_IN_URL, params={"key": FIREBASE_API_KEY}, json=payload)
        response_data = response.json()
        if 'email' in response_data:
            return {
                'email': response_data['email'],
                'username': response_data.get('displayName')
            }
        else:
            st.warning(response_data.get('error', {}).get('message', 'Signin failed'))
    except Exception as e:
        st.warning(f'Signin failed: {e}')

def reset_password(email):
    try:
        payload = {
            "email": email,
            "requestType": "PASSWORD_RESET"
        }
        response = requests.post(RESET_PASSWORD_URL, params={"key": FIREBASE_API_KEY}, json=payload)
        if response.status_code == 200:
            return True, "Reset email sent"
        else:
            error_message = response.json().get('error', {}).get('message', 'Password reset failed')
            return False, error_message
    except Exception as e:
        return False, str(e)

def authenticate_user():
    if 'username' not in st.session_state:
        st.session_state.username = ''
    if 'useremail' not in st.session_state:
        st.session_state.useremail = ''
    if 'signedout' not in st.session_state:
        st.session_state.signedout = False
    if 'signout' not in st.session_state:
        st.session_state.signout = False

    if not st.session_state.signedout:
        st.sidebar.header("Login / Signup")
        option = st.sidebar.selectbox("Choose an option", ["Login", "Sign Up", "Reset Password"])

        if option == "Login":
            st.sidebar.text_input("Email Address", key='login_email')
            st.sidebar.text_input("Password", type="password", key='login_password')
            if st.sidebar.button("Login"):
                user_info = sign_in(st.session_state.login_email, st.session_state.login_password)
                if user_info:
                    st.session_state.username = user_info['username']
                    st.session_state.useremail = user_info['email']
                    st.session_state.signedout = True
                    st.session_state.signout = True

        elif option == "Sign Up":
            st.sidebar.text_input("Email Address", key='signup_email')
            st.sidebar.text_input("Password", type="password", key='signup_password')
            st.sidebar.text_input("Username", key='signup_username')
            if st.sidebar.button("Sign Up"):
                user = sign_up(email=st.session_state.signup_email, password=st.session_state.signup_password, username=st.session_state.signup_username)
                if user:
                    st.success('Account created successfully!')
                    st.sidebar.text('Please Login using your email and password')

        elif option == "Reset Password":
            st.sidebar.text_input("Email Address", key='reset_email')
            if st.sidebar.button("Send Reset Link"):
                success, message = reset_password(st.session_state.reset_email)
                if success:
                    st.success(message)
                else:
                    st.warning(f"Password reset failed: {message}")

def show_authentication_page():
    st.title('Authentication Required')
    st.write("Please log in to access the application.")

def show_main_app():
    st.title("Main App Content")
    st.write("This is the main content of your app. Accessible only after login.")





# Initialize PubMed client
pubmed = PubMed(tool="Haystack2.0Prototype", email="dummyemail@gmail.com")

# Function to convert PubMed article to Haystack Document
def documentize(article):
    article_dict = article.toDict()
    pubmed_id = article_dict['pubmed_id'].partition('\n')[0]
    keywords = article_dict.get('keywords', [])
    return Document(
        content=article_dict['abstract'],
        meta={
            'pubmed_id': pubmed_id,
            'title': article_dict.get('title', ''),
            'keywords': keywords,
            'abstract': article_dict['abstract']
        }
    )

# Custom component to fetch articles from PubMed
@component
class PubMedFetcher():
    @component.output_types(articles=List[Document])
    def run(self, queries: List[str]):
        cleaned_queries = queries[0].strip().split('---')[0].strip().split('\n')[:3]
        print("Cleaned Queries:", cleaned_queries)

        articles = []
        try:
            query = f"({cleaned_queries[0]} AND {cleaned_queries[1]} AND {cleaned_queries[2]})"
            response = pubmed.query(query, max_results=5)
            for article in response:
                articles.append(documentize(article))
        except Exception as e:
            print(e)
            print(f"Couldn't fetch articles for queries: {queries}")
        for doc in articles:
            pubmed_id = doc.meta.get('pubmed_id', 'No pubmed_id')
            title = doc.meta.get('title', 'No title')
            st.markdown(f"PubMed ID: {pubmed_id}")
            st.markdown(f"Title: {title}")
        return {'articles': articles}

# Initialize HuggingFace generators
keyword_llm = HuggingFaceTGIGenerator(model="mistralai/Mixtral-8x7B-Instruct-v0.1")
keyword_llm.warm_up()
llm = HuggingFaceTGIGenerator(model="mistralai/Mixtral-8x7B-Instruct-v0.1")
llm.warm_up()

# Define prompt templates
keyword_prompt_template = """
Your task is to convert the following question into 3 keywords that can be used to find relevant medical research papers on PubMed.
Here is an example:
question: "What are the latest treatments for major depressive disorder?"
keywords:
Antidepressive Agents
Depressive Disorder, Major
Treatment-Resistant depression
---
question: {{ question }}
keywords:
"""

fmt_qa_prompt = """
Based on the given articles, provide a concise summary of

Articles:
{% for article in articles %}
Title: {{ article.meta['title'] }}
Abstract: {{ article.content }}
Keywords: {{ article.meta['keywords'] }}

{% endfor %}

Summary:
"""

# Initialize prompt builders
keyword_prompt_builder = PromptBuilder(template=keyword_prompt_template)
prompt_builder = PromptBuilder(template=fmt_qa_prompt)

# Initialize PubMed fetcher component
fetcher = PubMedFetcher()

# Build pipeline
pipe = Pipeline()
pipe.add_component("keyword_prompt_builder", keyword_prompt_builder)
pipe.add_component("keyword_llm", keyword_llm)
pipe.add_component("pubmed_fetcher", fetcher)
pipe.add_component("prompt_builder", prompt_builder)
pipe.add_component("llm", llm)

# Connect components in pipeline
pipe.connect("keyword_prompt_builder.prompt", "keyword_llm.prompt")
pipe.connect("keyword_llm.replies", "pubmed_fetcher.queries")
pipe.connect("pubmed_fetcher.articles", "prompt_builder.articles")
pipe.connect("prompt_builder.prompt", "llm.prompt")

# Function to ask a question and get a response
def ask_pubmed(question):
    output = pipe.run(data={
        "keyword_prompt_builder": {"question": question},
        "llm": {"generation_kwargs": {"max_new_tokens": 500}}
    })
    print("Question:", question)
    response = output['llm']['replies'][0]
    return response

# Set OpenAI API Key
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")
os.environ["LLAMA_CLOUD_API_KEY"] = "llx-zVZyYaZqw8H2Z2kXdFmCjO2fPWCWJQpATxUxJbNdr4q1xBFf"

# Initialize Database
def init_db():
    conn = sqlite3.connect('chat_history.db')
    cursor = conn.cursor()

    cursor.execute('''CREATE TABLE IF NOT EXISTS ChatSessions (
                      id INTEGER PRIMARY KEY AUTOINCREMENT,
                      start_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')

    cursor.execute('''CREATE TABLE IF NOT EXISTS Messages (
                      id INTEGER PRIMARY KEY AUTOINCREMENT,
                      session_id INTEGER,
                      content TEXT,
                      timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                      FOREIGN KEY (session_id) REFERENCES ChatSessions(id))''')

    conn.commit()
    conn.close()

init_db()

def use_llamaparse(file_content, file_name):
    with open(file_name, "wb") as f:
        f.write(file_content)

    parser = LlamaParse(result_type="markdown", verbose=True, language="en", num_workers=2)
    documents = parser.load_data([file_name])

    os.remove(file_name)

    res = ''
    for i in documents:
        res += i.text + " "
    return res

def save_session():
    conn = sqlite3.connect('chat_history.db')
    cursor = conn.cursor()
    cursor.execute('INSERT INTO ChatSessions DEFAULT VALUES')
    session_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return session_id

def save_message(session_id, message):
    conn = sqlite3.connect('chat_history.db')
    cursor = conn.cursor()
    cursor.execute('INSERT INTO Messages (session_id, content) VALUES (?, ?)', (session_id, message))
    conn.commit()
    conn.close()

def get_chat_sessions():
    conn = sqlite3.connect('chat_history.db')
    cursor = conn.cursor()
    cursor.execute('SELECT id, start_time FROM ChatSessions ORDER BY start_time DESC')
    sessions = cursor.fetchall()
    conn.close()
    return sessions

def get_messages(session_id):
    conn = sqlite3.connect('chat_history.db')
    cursor = conn.cursor()
    cursor.execute('SELECT content, timestamp FROM Messages WHERE session_id = ? ORDER BY timestamp ASC', (session_id,))
    messages = cursor.fetchall()
    conn.close()
    return messages

def generate_docx(result):
    from docx import Document
    doc = Document()
    doc.add_heading('Healthcare Diagnosis and Treatment Recommendations', 0)
    doc.add_paragraph(result)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio



def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Diagnosis and Treatment Plan</a>'

# Initialize Tools
search_tool = SerperDevTool()
scrape_tool = ScrapeWebsiteTool()

llm = ChatOpenAI(
    model="gpt-4o-mini",
    temperature=0.1,
    max_tokens=200  # Adjust the token limit as needed
)

# Define Agents
diagnostician = Agent(
    role="Medical Diagnostician",
    goal="Analyze patient symptoms and medical history to provide a preliminary diagnosis.",
    backstory="This agent specializes in diagnosing medical conditions based on patient-reported symptoms and medical history. It uses advanced algorithms and medical knowledge to identify potential health issues.",
    verbose=False,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

treatment_advisor = Agent(
    role="Treatment Advisor",
    goal="Recommend appropriate treatment plans based on the diagnosis provided by the Medical Diagnostician.",
    backstory="This agent specializes in creating treatment plans tailored to individual patient needs. It considers the diagnosis, patient history, and current best practices in medicine to recommend effective treatments.",
    verbose=False,
    allow_delegation=False,
    tools=[search_tool, scrape_tool],
    llm=llm
)

# Define Tasks
diagnose_task = Task(
    description=(
        "1. Analyze the patient's symptoms ({symptoms}) and medical history ({medical_history}).\n"
        "2. Provide a preliminary diagnosis with possible conditions based on the provided information.\n"
        "3. Limit the diagnosis to the most likely conditions."
    ),
    expected_output="A preliminary diagnosis with a list of possible conditions.",
    agent=diagnostician,
    timeout=20
)

treatment_task = Task(
    description=(
        "1. Based on the provided diagnosis, recommend appropriate treatment plans step by step.\n"
        "2. Consider the patient's medical history ({medical_history}) and current symptoms ({symptoms}).\n"
        "3. Provide detailed treatment recommendations, including medications, lifestyle changes, and follow-up care."
    ),
    expected_output="A comprehensive treatment plan tailored to the patient's needs.",
    agent=treatment_advisor,
    timeout=20
)

# Create Crew
crew = Crew(
    agents=[diagnostician, treatment_advisor],
    tasks=[diagnose_task, treatment_task],
    verbose=0
)

def diagnose_and_treat(symptoms, medical_history,session_id):

    diagnosis_input = {"symptoms": symptoms, "medical_history": medical_history}
    diagnosis_result = crew.kickoff(inputs=diagnosis_input)
                    # Display and summarize the diagnosis result
    st.write(diagnosis_result)
    diagnosis_summary = diagnosis_result['summary'] if 'summary' in diagnosis_result else diagnosis_result

                    # Save diagnosis result to chat history
    save_message(session_id, diagnosis_summary)

                    # Prepare input for treatment task
    treatment_input = {"symptoms": symptoms, "medical_history": medical_history, "diagnosis": diagnosis_summary}
    treatment_result = crew.kickoff(inputs=treatment_input)

                    # Display and summarize the treatment result
    st.write(treatment_result)
    treatment_summary = treatment_result['summary'] if 'summary' in treatment_result else treatment_result

                    # Save treatment result to chat history
    save_message(session_id, treatment_summary)

                    # Combine results
    final_result = f"Diagnosis: {diagnosis_summary}\n\nTreatment Plan: {treatment_summary}"
    return final_result

class MultiApp:
    def __init__(self):
        self.apps = []

    def add_app(self, title, func):
        self.apps.append({
            "title": title,
            "function": func
        })
    def run(self):
        authenticate_user()

        if st.session_state.signedout:
          with st.sidebar:
            app = option_menu(
                menu_title='Pondering',
                options=['Home', 'PubMed Query', 'Diagnosis & Treatment'],
                icons=['house-fill', 'search', 'chat-fill'],
                menu_icon='menu-button',
                default_index=0,
                styles={
                     "container": {"padding": "5!important", "background-color": '#fafafa'},
                    "icon": {"color": "black", "font-size": "25px"},
                    "nav-link": {"font-size": "18px", "text-align": "left", "margin": "0px", "--hover-color": "#eee"},
                    "nav-link-selected": {"background-color": "#02ab21"},
                }
            )

          if app == "Home":
            # Sidebar for chat history
            st.sidebar.title("Chat History")
            chat_sessions = get_chat_sessions()
            for session in chat_sessions:
                if st.sidebar.button(f"Session {session[0]} - {session[1]}"):
                    messages = get_messages(session[0])
                    st.sidebar.write(f"Session {session[0]} - {session[1]}")
                    for message in messages:
                        st.sidebar.write(f"{message[1]}: {message[0]}")
            st.title("Welcome to CURA.AI")
            st.write("Hello there, how can I help you?")

            # Chatbot code
            # Initialize the model
            model = Ollama(model="mistral")

            # Define the prompt template
            prompt = ChatPromptTemplate.from_messages(
                [
                    ("system", "You're a very knowledgeable doctor who provides accurate  answers to medical related questions and You must not answer question out of medical related topics."),
                    ("human", "{question}")
                ]
            )

            # Create the runnable pipeline
            runnable = prompt | model | StrOutputParser()

            # Input form for user question
            with st.form(key="question_form"):
                user_question = st.text_input("Ask me a question:")
                submit_button = st.form_submit_button(label="Ask")

            # Process user input and display the response
            if submit_button and user_question:
                with st.spinner("Thinking..."):
                    # Process the question using the runnable pipeline
                    result = runnable.invoke({"question": user_question})

                    # Display the response
                    st.write("LLM's Response:")
                    st.write(result)

            # Provide a text box for continuous conversation
            if 'responses' not in st.session_state:
                st.session_state['responses'] = []

            if 'user_input' not in st.session_state:
                st.session_state['user_input'] = ''

            # Async function to get response from the model
            async def get_response_async(user_input):
                loop = asyncio.get_running_loop()
                with concurrent.futures.ThreadPoolExecutor() as pool:
                    result = await loop.run_in_executor(pool, runnable.invoke, {"question": user_input})
                return result

            def get_response():
                user_input = st.session_state.user_input
                if user_input:
                    with st.spinner("Thinking..."):
                        result = asyncio.run(get_response_async(user_input))
                        st.session_state.responses.append(("You", user_input))
                        st.session_state.responses.append(("LLM", result))
                    st.session_state.user_input = ''

            st.text_input("Continue the conversation:", key='user_input', on_change=get_response)

            # Display conversation history
            for sender, message in st.session_state['responses']:
                st.write(f"**{sender}:** {message}")

          elif app == "PubMed Query":
            st.title("PubMed Article Query")
            question = st.text_area("Enter your research question:")
            if st.button("Fetch Articles"):
                with st.spinner('Fetching articles...'):
                    result = ask_pubmed(question)
                    st.write(result)

          elif app == "Diagnosis & Treatment":
            st.title("Healthcare Diagnosis and Treatment Recommendations")
            symptoms = st.text_area('Enter Symptoms', 'e.g., fever, cough, headache')
            medical_history = st.text_area('Enter Medical History', 'e.g., diabetes, hypertension')
            uploaded_files = st.file_uploader("Upload Documents", type=["pdf", "csv", "xlsx", "docx", "pptx"], accept_multiple_files=True)
            if uploaded_files:
                for uploaded_file in uploaded_files:
                     file_content = uploaded_file.read()
                     temp_file_name = uploaded_file.name
                     parsed_text = use_llamaparse(file_content, temp_file_name)
                     medical_history += "\n" + parsed_text  # Append parsed text to medical history
            if st.button("Get Diagnosis and Treatment Plan"):
                with st.spinner('Generating recommendations...'):
                    # Start a new chat session
                    session_id = save_session()





                    # Combine results
                    final_result = diagnose_and_treat(symptoms, medical_history,session_id)

                    docx_file = generate_docx(final_result)

                    download_link = get_download_link(docx_file, "diagnosis_and_treatment_plan.docx")

                    st.markdown(download_link, unsafe_allow_html=True)
        else:
          show_authentication_page()


# Run the app
app = MultiApp()
app.add_app("Home", lambda: st.write("Welcome to the AI Medical Assistant"))
app.add_app("PubMed Query", lambda: ask_pubmed(st.text_area("Enter your research question:")))
app.add_app("Diagnosis & Treatment", diagnose_and_treat)
app.run()
