# import streamlit as st
# from crewai import Agent, Task, Crew, Process
# import os
# from crewai_tools import ScrapeWebsiteTool, SerperDevTool
# from dotenv import load_dotenv
# from langchain_openai import ChatOpenAI
# from docx import Document
# from io import BytesIO
# import base64
# import time

# load_dotenv()

# # LLM object and API Key
# os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
# os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")


# def generate_docx(result):
#     doc = Document()
#     doc.add_heading('Healthcare Diagnosis and Treatment Recommendations', 0)
#     doc.add_paragraph(result)
#     bio = BytesIO()
#     doc.save(bio)
#     bio.seek(0)
#     return bio

# def get_download_link(bio, filename):
#     b64 = base64.b64encode(bio.read()).decode()
#     return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Diagnosis and Treatment Plan</a>'

# st.set_page_config(
#     layout="wide"
# )

# # Title
# st.title("AI Agents to Empower Doctors")

# # Text Inputs
# gender = st.selectbox('Select Gender', ('Male', 'Female', 'Other'))
# age = st.number_input('Enter Age', min_value=0, max_value=120, value=25)
# symptoms = st.text_area('Enter Symptoms', 'e.g., fever, cough, headache')
# medical_history = st.text_area('Enter Medical History', 'e.g., diabetes, hypertension')

# # Initialize Tools
# search_tool = SerperDevTool()
# scrape_tool = ScrapeWebsiteTool()

# llm = ChatOpenAI(
#     model="gpt-3.5-turbo",
#     temperature=0.1,
#     max_tokens=200  # Adjust the token limit as needed
# )

# # Define Agents
# diagnostician = Agent(
#     role="Medical Diagnostician",
#     goal="Analyze patient symptoms and medical history to provide a preliminary diagnosis.",
#     backstory="This agent specializes in diagnosing medical conditions based on patient-reported symptoms and medical history. It uses advanced algorithms and medical knowledge to identify potential health issues.",
#     verbose=False,  # Reduce verbosity
#     allow_delegation=False,
#     tools=[search_tool, scrape_tool],
#     llm=llm
# )

# treatment_advisor = Agent(
#     role="Treatment Advisor",
#     goal="Recommend appropriate treatment plans based on the diagnosis provided by the Medical Diagnostician.",
#     backstory="This agent specializes in creating treatment plans tailored to individual patient needs. It considers the diagnosis, patient history, and current best practices in medicine to recommend effective treatments.",
#     verbose=False,  # Reduce verbosity
#     allow_delegation=False,
#     tools=[search_tool, scrape_tool],
#     llm=llm
# )

# # Define Tasks
# diagnose_task = Task(
#     description=(
#         "1. Analyze the patient's symptoms ({symptoms}) and medical history ({medical_history}).\n"
#         "2. Provide a preliminary diagnosis with possible conditions based on the provided information.\n"
#         "3. Limit the diagnosis to the most likely conditions."
#     ),
#     expected_output="A preliminary diagnosis with a list of possible conditions.",
#     agent=diagnostician,
#     timeout=20  # Timeout in seconds
# )

# treatment_task = Task(
#     description=(
#         "1. Based on the provided diagnosis, recommend appropriate treatment plans step by step.\n"
#         "2. Consider the patient's medical history ({medical_history}) and current symptoms ({symptoms}).\n"
#         "3. Provide detailed treatment recommendations, including medications, lifestyle changes, and follow-up care."
#     ),
#     expected_output="A comprehensive treatment plan tailored to the patient's needs.",
#     agent=treatment_advisor,
#     timeout=20  # Timeout in seconds
# )

# # Create Crew
# crew = Crew(
#     agents=[diagnostician, treatment_advisor],
#     tasks=[diagnose_task, treatment_task],
#     verbose=0
# )

# # Execution
# if st.button("Get Diagnosis and Treatment Plan"):
#     with st.spinner('Generating recommendations...'):
#         diagnosis_input = {"symptoms": symptoms, "medical_history": medical_history}
#         diagnosis_result = crew.kickoff(inputs=diagnosis_input)
        
#         # Display and summarize the diagnosis result
#         st.write(diagnosis_result)
#         diagnosis_summary = diagnosis_result['summary'] if 'summary' in diagnosis_result else diagnosis_result
        
#         # Prepare input for treatment task
#         treatment_input = {"symptoms": symptoms, "medical_history": medical_history, "diagnosis": diagnosis_summary}
#         treatment_result = crew.kickoff(inputs=treatment_input)
        
#         # Display and summarize the treatment result
#         st.write(treatment_result)
#         treatment_summary = treatment_result['summary'] if 'summary' in treatment_result else treatment_result
        
#         # Combine results
#         final_result = f"Diagnosis: {diagnosis_summary}\n\nTreatment Plan: {treatment_summary}"
#         docx_file = generate_docx(final_result)

#         download_link = get_download_link(docx_file, "diagnosis_and_treatment_plan.docx")
        
#         st.markdown(download_link, unsafe_allow_html=True)




# import sqlite3

# def init_db():
#     conn = sqlite3.connect('chat_history.db')
#     cursor = conn.cursor()
    
#     cursor.execute('''CREATE TABLE IF NOT EXISTS ChatSessions (
#                       id INTEGER PRIMARY KEY AUTOINCREMENT,
#                       start_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
                      
#     cursor.execute('''CREATE TABLE IF NOT EXISTS Messages (
#                       id INTEGER PRIMARY KEY AUTOINCREMENT,
#                       session_id INTEGER,
#                       content TEXT,
#                       timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#                       FOREIGN KEY (session_id) REFERENCES ChatSessions(id))''')
                      
#     conn.commit()
#     conn.close()

# init_db()
# def save_session():
#     conn = sqlite3.connect('chat_history.db')
#     cursor = conn.cursor()
#     cursor.execute('INSERT INTO ChatSessions DEFAULT VALUES')
#     session_id = cursor.lastrowid
#     conn.commit()
#     conn.close()
#     return session_id

# def save_message(session_id, message):
#     conn = sqlite3.connect('chat_history.db')
#     cursor = conn.cursor()
#     cursor.execute('INSERT INTO Messages (session_id, content) VALUES (?, ?)', (session_id, message))
#     conn.commit()
#     conn.close()

# def get_chat_sessions():
#     conn = sqlite3.connect('chat_history.db')
#     cursor = conn.cursor()
#     cursor.execute('SELECT id, start_time FROM ChatSessions ORDER BY start_time DESC')
#     sessions = cursor.fetchall()
#     conn.close()
#     return sessions

# def get_messages(session_id):
#     conn = sqlite3.connect('chat_history.db')
#     cursor = conn.cursor()
#     cursor.execute('SELECT content, timestamp FROM Messages WHERE session_id = ? ORDER BY timestamp ASC', (session_id,))
#     messages = cursor.fetchall()
#     conn.close()
#     return messages
# import streamlit as st
# from crewai import Agent, Task, Crew
# from crewai_tools import ScrapeWebsiteTool, SerperDevTool
# from dotenv import load_dotenv
# from langchain_openai import ChatOpenAI
# from docx import Document
# from io import BytesIO
# import base64
# import sqlite3
# import os

# load_dotenv()

# # LLM object and API Key
# os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
# os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")

# # Initialize Database
# init_db()

# def generate_docx(result):
#     doc = Document()
#     doc.add_heading('Healthcare Diagnosis and Treatment Recommendations', 0)
#     doc.add_paragraph(result)
#     bio = BytesIO()
#     doc.save(bio)
#     bio.seek(0)
#     return bio

# def get_download_link(bio, filename):
#     b64 = base64.b64encode(bio.read()).decode()
#     return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Diagnosis and Treatment Plan</a>'

# st.set_page_config(
#     layout="wide"
# )

# # Title
# st.title("AI Agents to Empower Doctors")

# # Text Inputs
# gender = st.selectbox('Select Gender', ('Male', 'Female', 'Other'))
# age = st.number_input('Enter Age', min_value=0, max_value=120, value=25)
# symptoms = st.text_area('Enter Symptoms', 'e.g., fever, cough, headache')
# medical_history = st.text_area('Enter Medical History', 'e.g., diabetes, hypertension')

# # Initialize Tools
# search_tool = SerperDevTool()
# scrape_tool = ScrapeWebsiteTool()

# llm = ChatOpenAI(
#     model="gpt-3.5-turbo",
#     temperature=0.1,
#     max_tokens=200  # Adjust the token limit as needed
# )

# # Define Agents
# diagnostician = Agent(
#     role="Medical Diagnostician",
#     goal="Analyze patient symptoms and medical history to provide a preliminary diagnosis.",
#     backstory="This agent specializes in diagnosing medical conditions based on patient-reported symptoms and medical history. It uses advanced algorithms and medical knowledge to identify potential health issues.",
#     verbose=False,  # Reduce verbosity
#     allow_delegation=False,
#     tools=[search_tool, scrape_tool],
#     llm=llm
# )

# treatment_advisor = Agent(
#     role="Treatment Advisor",
#     goal="Recommend appropriate treatment plans based on the diagnosis provided by the Medical Diagnostician.",
#     backstory="This agent specializes in creating treatment plans tailored to individual patient needs. It considers the diagnosis, patient history, and current best practices in medicine to recommend effective treatments.",
#     verbose=False,  # Reduce verbosity
#     allow_delegation=False,
#     tools=[search_tool, scrape_tool],
#     llm=llm
# )

# # Define Tasks
# diagnose_task = Task(
#     description=(
#         "1. Analyze the patient's symptoms ({symptoms}) and medical history ({medical_history}).\n"
#         "2. Provide a preliminary diagnosis with possible conditions based on the provided information.\n"
#         "3. Limit the diagnosis to the most likely conditions."
#     ),
#     expected_output="A preliminary diagnosis with a list of possible conditions.",
#     agent=diagnostician,
#     timeout=20  # Timeout in seconds
# )

# treatment_task = Task(
#     description=(
#         "1. Based on the provided diagnosis, recommend appropriate treatment plans step by step.\n"
#         "2. Consider the patient's medical history ({medical_history}) and current symptoms ({symptoms}).\n"
#         "3. Provide detailed treatment recommendations, including medications, lifestyle changes, and follow-up care."
#     ),
#     expected_output="A comprehensive treatment plan tailored to the patient's needs.",
#     agent=treatment_advisor,
#     timeout=20  # Timeout in seconds
# )

# # Create Crew
# crew = Crew(
#     agents=[diagnostician, treatment_advisor],
#     tasks=[diagnose_task, treatment_task],
#     verbose=0
# )

# # Execution
# if st.button("Get Diagnosis and Treatment Plan"):
#     with st.spinner('Generating recommendations...'):
#         # Start a new chat session
#         session_id = save_session()
        
#         diagnosis_input = {"symptoms": symptoms, "medical_history": medical_history}
#         diagnosis_result = crew.kickoff(inputs=diagnosis_input)
        
#         # Display and summarize the diagnosis result
#         st.write(diagnosis_result)
#         diagnosis_summary = diagnosis_result['summary'] if 'summary' in diagnosis_result else diagnosis_result
        
#         # Save diagnosis result to chat history
#         save_message(session_id, diagnosis_summary)
        
#         # Prepare input for treatment task
#         treatment_input = {"symptoms": symptoms, "medical_history": medical_history, "diagnosis": diagnosis_summary}
#         treatment_result = crew.kickoff(inputs=treatment_input)
        
#         # Display and summarize the treatment result
#         st.write(treatment_result)
#         treatment_summary = treatment_result['summary'] if 'summary' in treatment_result else treatment_result
        
#         # Save treatment result to chat history
#         save_message(session_id, treatment_summary)
        
#         # Combine results
#         final_result = f"Diagnosis: {diagnosis_summary}\n\nTreatment Plan: {treatment_summary}"
#         docx_file = generate_docx(final_result)

#         download_link = get_download_link(docx_file, "diagnosis_and_treatment_plan.docx")
        
#         st.markdown(download_link, unsafe_allow_html=True)

# # Chat History Display
# st.sidebar.title("Chat History")

# chat_sessions = get_chat_sessions()
# for session in chat_sessions:
#     if st.sidebar.button(f"Session {session[0]} - {session[1]}"):
#         messages = get_messages(session[0])
#         st.sidebar.write(f"Session {session[0]} - {session[1]}")
#         for message in messages:
#             st.sidebar.write(f"{message[1]}: {message[0]}")




# import streamlit as st
# from crewai import Agent, Task, Crew
# from crewai_tools import ScrapeWebsiteTool, SerperDevTool
# from dotenv import load_dotenv
# from langchain_openai import ChatOpenAI
# from docx import Document
# from io import BytesIO
# from llama_parse import LlamaParse
# import base64
# import sqlite3
# import os

# load_dotenv()

# # LLM object and API Key
# os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
# os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")
# os.environ["LLAMA_CLOUD_API_KEY"] = "llx-zVZyYaZqw8H2Z2kXdFmCjO2fPWCWJQpATxUxJbNdr4q1xBFf"

# # Initialize Database
# def init_db():
#     conn = sqlite3.connect('chat_history.db')
#     cursor = conn.cursor()
    
#     cursor.execute('''CREATE TABLE IF NOT EXISTS ChatSessions (
#                       id INTEGER PRIMARY KEY AUTOINCREMENT,
#                       start_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
                      
#     cursor.execute('''CREATE TABLE IF NOT EXISTS Messages (
#                       id INTEGER PRIMARY KEY AUTOINCREMENT,
#                       session_id INTEGER,
#                       content TEXT,
#                       timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
#                       FOREIGN KEY (session_id) REFERENCES ChatSessions(id))''')
                      
#     conn.commit()
#     conn.close()

# init_db()

# def use_llamaparse(file_content, file_name):
#     with open(file_name, "wb") as f:
#         f.write(file_content)
    
#     parser = LlamaParse(result_type="markdown", verbose=True, language="en", num_workers=2)
#     documents = parser.load_data([file_name])
    
#     os.remove(file_name)
    
#     res = ''
#     for i in documents:
#         res += i.text + " "
#     return res    

# def save_session():
#     conn = sqlite3.connect('chat_history.db')
#     cursor = conn.cursor()
#     cursor.execute('INSERT INTO ChatSessions DEFAULT VALUES')
#     session_id = cursor.lastrowid
#     conn.commit()
#     conn.close()
#     return session_id

# def save_message(session_id, message):
#     conn = sqlite3.connect('chat_history.db')
#     cursor = conn.cursor()
#     cursor.execute('INSERT INTO Messages (session_id, content) VALUES (?, ?)', (session_id, message))
#     conn.commit()
#     conn.close()

# def get_chat_sessions():
#     conn = sqlite3.connect('chat_history.db')
#     cursor = conn.cursor()
#     cursor.execute('SELECT id, start_time FROM ChatSessions ORDER BY start_time DESC')
#     sessions = cursor.fetchall()
#     conn.close()
#     return sessions

# def get_messages(session_id):
#     conn = sqlite3.connect('chat_history.db')
#     cursor = conn.cursor()
#     cursor.execute('SELECT content, timestamp FROM Messages WHERE session_id = ? ORDER BY timestamp ASC', (session_id,))
#     messages = cursor.fetchall()
#     conn.close()
#     return messages

# def generate_docx(result):
#     doc = Document()
#     doc.add_heading('Healthcare Diagnosis and Treatment Recommendations', 0)
#     doc.add_paragraph(result)
#     bio = BytesIO()
#     doc.save(bio)
#     bio.seek(0)
#     return bio

# def get_download_link(bio, filename):
#     b64 = base64.b64encode(bio.read()).decode()
#     return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Diagnosis and Treatment Plan</a>'

# st.set_page_config(
#     layout="wide"
# )

# # Title
# st.title("AI Agents to Empower Doctors")

# # Create columns
# main_col, sidebar_col = st.columns([3, 1])

# # Main content in main_col
# with main_col:
#     # Text Inputs
#     gender = st.selectbox('Select Gender', ('Male', 'Female', 'Other'))
#     age = st.number_input('Enter Age', min_value=0, max_value=120, value=25)
#     symptoms = st.text_area('Enter Symptoms', 'e.g., fever, cough, headache')

#     # Medical History Text Area with Document Upload
#     medical_history = st.text_area('Enter Medical History', 'e.g., diabetes, hypertension')

#     uploaded_files = st.file_uploader("Upload Documents", type=["pdf", "csv", "xlsx", "docx", "pptx"], accept_multiple_files=True)
#     if uploaded_files:
#         for uploaded_file in uploaded_files:
#             file_content = uploaded_file.read()
#             temp_file_name = uploaded_file.name
#             print(temp_file_name, ".....", file_content)
#             parsed_text = use_llamaparse(file_content, temp_file_name)
#             medical_history += "\n" + parsed_text  # Append parsed text to medical history

#     # Initialize Tools
#     search_tool = SerperDevTool()
#     scrape_tool = ScrapeWebsiteTool()

#     llm = ChatOpenAI(
#         model="gpt-3.5-turbo",
#         temperature=0.1,
#         max_tokens=200  # Adjust the token limit as needed
#     )

#     # Define Agents
#     diagnostician = Agent(
#         role="Medical Diagnostician",
#         goal="Analyze patient symptoms and medical history to provide a preliminary diagnosis.",
#         backstory="This agent specializes in diagnosing medical conditions based on patient-reported symptoms and medical history. It uses advanced algorithms and medical knowledge to identify potential health issues.",
#         verbose=False,  # Reduce verbosity
#         allow_delegation=False,
#         tools=[search_tool, scrape_tool],
#         llm=llm
#     )

#     treatment_advisor = Agent(
#         role="Treatment Advisor",
#         goal="Recommend appropriate treatment plans based on the diagnosis provided by the Medical Diagnostician.",
#         backstory="This agent specializes in creating treatment plans tailored to individual patient needs. It considers the diagnosis, patient history, and current best practices in medicine to recommend effective treatments.",
#         verbose=False,  # Reduce verbosity
#         allow_delegation=False,
#         tools=[search_tool, scrape_tool],
#         llm=llm
#     )

#     # Define Tasks
#     diagnose_task = Task(
#         description=(
#             "1. Analyze the patient's symptoms ({symptoms}) and medical history ({medical_history}).\n"
#             "2. Provide a preliminary diagnosis with possible conditions based on the provided information.\n"
#             "3. Limit the diagnosis to the most likely conditions."
#         ),
#         expected_output="A preliminary diagnosis with a list of possible conditions.",
#         agent=diagnostician,
#         timeout=20  # Timeout in seconds
#     )

#     treatment_task = Task(
#         description=(
#             "1. Based on the provided diagnosis, recommend appropriate treatment plans step by step.\n"
#             "2. Consider the patient's medical history ({medical_history}) and current symptoms ({symptoms}).\n"
#             "3. Provide detailed treatment recommendations, including medications, lifestyle changes, and follow-up care."
#         ),
#         expected_output="A comprehensive treatment plan tailored to the patient's needs.",
#         agent=treatment_advisor,
#         timeout=20  # Timeout in seconds
#     )

#     # Create Crew
#     crew = Crew(
#         agents=[diagnostician, treatment_advisor],
#         tasks=[diagnose_task, treatment_task],
#         verbose=0
#     )

#     # Execution
#     if st.button("Get Diagnosis and Treatment Plan"):
#         with st.spinner('Generating recommendations...'):
#             # Start a new chat session
#             session_id = save_session()
            
#             diagnosis_input = {"symptoms": symptoms, "medical_history": medical_history}
#             diagnosis_result = crew.kickoff(inputs=diagnosis_input)
            
#             # Display and summarize the diagnosis result
#             st.write(diagnosis_result)
#             diagnosis_summary = diagnosis_result['summary'] if 'summary' in diagnosis_result else diagnosis_result
            
#             # Save diagnosis result to chat history
#             save_message(session_id, diagnosis_summary)
            
#             # Prepare input for treatment task
#             treatment_input = {"symptoms": symptoms, "medical_history": medical_history, "diagnosis": diagnosis_summary}
#             treatment_result = crew.kickoff(inputs=treatment_input)
            
#             # Display and summarize the treatment result
#             st.write(treatment_result)
#             treatment_summary = treatment_result['summary'] if 'summary' in treatment_result else treatment_result
            
#             # Save treatment result to chat history
#             save_message(session_id, treatment_summary)
            
#             # Combine results
#             final_result = f"Diagnosis: {diagnosis_summary}\n\nTreatment Plan: {treatment_summary}"
#             docx_file = generate_docx(final_result)

#             download_link = get_download_link(docx_file, "diagnosis_and_treatment_plan.docx")
            
#             st.markdown(download_link, unsafe_allow_html=True)

# # Chat History Display in sidebar_col
# with sidebar_col:
#     st.title("Chat History")

#     chat_sessions = get_chat_sessions()
#     for session in chat_sessions:
#         if st.button(f"Session {session[0]} - {session[1]}"):
#             messages = get_messages(session[0])
#             st.write(f"Session {session[0]} - {session[1]}")
#             for message in messages:
#                 st.write(f"{message[1]}: {message[0]}")


import streamlit as st
from crewai import Agent, Task, Crew
from crewai_tools import ScrapeWebsiteTool, SerperDevTool
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from docx import Document
from io import BytesIO
from llama_parse import LlamaParse
import base64
import sqlite3
import os

load_dotenv()

# LLM object and API Key
os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")
os.environ["SERPER_API_KEY"] = os.getenv("SERPER_API_KEY")
os.environ["LLAMA_CLOUD_API_KEY"] = "llx-zVZyYaZqw8H2Z2kXdFmCjO2fPWCWJQpATxUxJbNdr4q1xBFf"

# Initialize Database
def init_db():
    conn = sqlite3.connect('chat_history.db')
    cursor = conn.cursor()
    
    cursor.execute('''CREATE TABLE IF NOT EXISTS ChatSessions (
                      id INTEGER PRIMARY KEY AUTOINCREMENT,
                      start_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
                      
    cursor.execute('''CREATE TABLE IF NOT EXISTS Messages (
                      id INTEGER PRIMARY KEY AUTOINCREMENT,
                      session_id INTEGER,
                      content TEXT,
                      timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                      FOREIGN KEY (session_id) REFERENCES ChatSessions(id))''')
                      
    conn.commit()
    conn.close()

init_db()

def use_llamaparse(file_content, file_name):
    with open(file_name, "wb") as f:
        f.write(file_content)
    
    parser = LlamaParse(result_type="markdown", verbose=True, language="en", num_workers=2)
    documents = parser.load_data([file_name])
    
    os.remove(file_name)
    
    res = ''
    for i in documents:
        res += i.text + " "
    return res    

def save_session():
    conn = sqlite3.connect('chat_history.db')
    cursor = conn.cursor()
    cursor.execute('INSERT INTO ChatSessions DEFAULT VALUES')
    session_id = cursor.lastrowid
    conn.commit()
    conn.close()
    return session_id

def save_message(session_id, message):
    conn = sqlite3.connect('chat_history.db')
    cursor = conn.cursor()
    cursor.execute('INSERT INTO Messages (session_id, content) VALUES (?, ?)', (session_id, message))
    conn.commit()
    conn.close()

def get_chat_sessions():
    conn = sqlite3.connect('chat_history.db')
    cursor = conn.cursor()
    cursor.execute('SELECT id, start_time FROM ChatSessions ORDER BY start_time DESC')
    sessions = cursor.fetchall()
    conn.close()
    return sessions

def get_messages(session_id):
    conn = sqlite3.connect('chat_history.db')
    cursor = conn.cursor()
    cursor.execute('SELECT content, timestamp FROM Messages WHERE session_id = ? ORDER BY timestamp ASC', (session_id,))
    messages = cursor.fetchall()
    conn.close()
    return messages

def generate_docx(result):
    doc = Document()
    doc.add_heading('Healthcare Diagnosis and Treatment Recommendations', 0)
    doc.add_paragraph(result)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Diagnosis and Treatment Plan</a>'

st.set_page_config(
    layout="wide"
)

# Title
st.title("AI Agents to Empower Doctors")

# Sidebar for chat history
st.sidebar.title("Chat History")
chat_sessions = get_chat_sessions()
for session in chat_sessions:
    if st.sidebar.button(f"Session {session[0]} - {session[1]}"):
        messages = get_messages(session[0])
        st.sidebar.write(f"Session {session[0]} - {session[1]}")
        for message in messages:
            st.sidebar.write(f"{message[1]}: {message[0]}")

# Create main columns
main_col = st.container()

# Main content in main_col
with main_col:
    # Text Inputs
    gender = st.selectbox('Select Gender', ('Male', 'Female', 'Other'))
    age = st.number_input('Enter Age', min_value=0, max_value=120, value=25)
    symptoms = st.text_area('Enter Symptoms', 'e.g., fever, cough, headache')

    # Medical History Text Area with Document Upload
    medical_history = st.text_area('Enter Medical History', 'e.g., diabetes, hypertension')

    uploaded_files = st.file_uploader("Upload Documents", type=["pdf", "csv", "xlsx", "docx", "pptx"], accept_multiple_files=True)
    if uploaded_files:
        for uploaded_file in uploaded_files:
            file_content = uploaded_file.read()
            temp_file_name = uploaded_file.name
            parsed_text = use_llamaparse(file_content, temp_file_name)
            medical_history += "\n" + parsed_text  # Append parsed text to medical history

    # Initialize Tools
    search_tool = SerperDevTool()
    scrape_tool = ScrapeWebsiteTool()

    llm = ChatOpenAI(
        model="gpt-3.5-turbo",
        temperature=0.1,
        max_tokens=200  # Adjust the token limit as needed
    )

    # Define Agents
    diagnostician = Agent(
        role="Medical Diagnostician",
        goal="Analyze patient symptoms and medical history to provide a preliminary diagnosis.",
        backstory="This agent specializes in diagnosing medical conditions based on patient-reported symptoms and medical history. It uses advanced algorithms and medical knowledge to identify potential health issues.",
        verbose=False,  # Reduce verbosity
        allow_delegation=False,
        tools=[search_tool, scrape_tool],
        llm=llm
    )

    treatment_advisor = Agent(
        role="Treatment Advisor",
        goal="Recommend appropriate treatment plans based on the diagnosis provided by the Medical Diagnostician.",
        backstory="This agent specializes in creating treatment plans tailored to individual patient needs. It considers the diagnosis, patient history, and current best practices in medicine to recommend effective treatments.",
        verbose=False,  # Reduce verbosity
        allow_delegation=False,
        tools=[search_tool, scrape_tool],
        llm=llm
    )

    # Define Tasks
    diagnose_task = Task(
        description=(
            "1. Analyze the patient's symptoms ({symptoms}) and medical history ({medical_history}).\n"
            "2. Provide a preliminary diagnosis with possible conditions based on the provided information.\n"
            "3. Limit the diagnosis to the most likely conditions."
        ),
        expected_output="A preliminary diagnosis with a list of possible conditions.",
        agent=diagnostician,
        timeout=20  # Timeout in seconds
    )

    treatment_task = Task(
        description=(
            "1. Based on the provided diagnosis, recommend appropriate treatment plans step by step.\n"
            "2. Consider the patient's medical history ({medical_history}) and current symptoms ({symptoms}).\n"
            "3. Provide detailed treatment recommendations, including medications, lifestyle changes, and follow-up care."
        ),
        expected_output="A comprehensive treatment plan tailored to the patient's needs.",
        agent=treatment_advisor,
        timeout=20  # Timeout in seconds
    )

    # Create Crew
    crew = Crew(
        agents=[diagnostician, treatment_advisor],
        tasks=[diagnose_task, treatment_task],
        verbose=0
    )

    # Execution
    if st.button("Get Diagnosis and Treatment Plan"):
        with st.spinner('Generating recommendations...'):
            # Start a new chat session
            session_id = save_session()
            
            diagnosis_input = {"symptoms": symptoms, "medical_history": medical_history}
            diagnosis_result = crew.kickoff(inputs=diagnosis_input)
            
            # Display and summarize the diagnosis result
            st.write(diagnosis_result)
            diagnosis_summary = diagnosis_result['summary'] if 'summary' in diagnosis_result else diagnosis_result
            
            # Save diagnosis result to chat history
            save_message(session_id, diagnosis_summary)
            
            # Prepare input for treatment task
            treatment_input = {"symptoms": symptoms, "medical_history": medical_history, "diagnosis": diagnosis_summary}
            treatment_result = crew.kickoff(inputs=treatment_input)
            
            # Display and summarize the treatment result
            st.write(treatment_result)
            treatment_summary = treatment_result['summary'] if 'summary' in treatment_result else treatment_result
            
            # Save treatment result to chat history
            save_message(session_id, treatment_summary)
            
            # Combine results
            final_result = f"Diagnosis: {diagnosis_summary}\n\nTreatment Plan: {treatment_summary}"
            docx_file = generate_docx(final_result)

            download_link = get_download_link(docx_file, "diagnosis_and_treatment_plan.docx")
            
            st.markdown(download_link, unsafe_allow_html=True)
